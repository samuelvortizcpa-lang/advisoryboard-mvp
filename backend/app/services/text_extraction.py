"""
Text extraction from uploaded documents.

Supported types:
  pdf        → pdfplumber (handles scanned-light PDFs, tables, columns)
               with automatic Tesseract OCR fallback for garbled output
  docx       → python-docx (paragraphs + tables)
  doc        → falls back to docx attempt, raises clear error if it fails
  txt        → plain UTF-8 read  (Fathom .txt exports work out of the box)
  csv        → rows joined as readable text
  json       → Fathom meeting transcript JSON export (title, date, attendees,
                transcript entries, action items)
  eml        → Python's built-in email library (RFC 2822 standard email)
  msg        → extract-msg library (Outlook/MAPI format)
  mp4/mov    → ffmpeg extracts audio → OpenAI Whisper transcription
  mp3/m4a/wav→ OpenAI Whisper transcription (chunked if > 24 MB)
  xlsx / xls / pptx → raises UnsupportedFileType so caller can mark gracefully
"""

from __future__ import annotations

import csv
import io
import logging
from pathlib import Path

logger = logging.getLogger(__name__)

# Audio/video types are handled by the Whisper transcription service rather
# than a local text parser.
_AUDIO_TYPES: frozenset[str] = frozenset({"mp4", "m4a", "mp3", "wav"})


class ExtractionError(Exception):
    """Raised when text cannot be extracted from a file."""


class UnsupportedFileType(ExtractionError):
    """Raised for file types that cannot be processed into text."""


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------


def extract_text(file_path: str, file_type: str) -> str:
    """
    Extract raw text from *file_path*.

    Returns a (possibly empty) string.  Raises ExtractionError on hard failures.

    Audio/video files are transcribed via OpenAI Whisper (requires ffmpeg for
    video files and for audio files larger than 24 MB).
    """
    path = Path(file_path)
    if not path.exists():
        raise ExtractionError(f"File not found: {file_path}")

    ext = file_type.lower().lstrip(".")

    # --- audio/video: delegate to Whisper transcription service ---
    if ext in _AUDIO_TYPES:
        from app.services.audio_transcriber import transcribe_audio
        try:
            text = transcribe_audio(file_path, ext)
        except (RuntimeError, ValueError) as exc:
            raise ExtractionError(str(exc)) from exc
    else:
        # --- document types: local text parsers ---
        extractors = {
            "pdf": _extract_pdf,
            "docx": _extract_docx,
            "doc": _extract_docx,   # may fail for old binary .doc
            "txt": _extract_txt,
            "csv": _extract_csv,
            "json": _extract_fathom_json,
            "eml": _extract_email,
            "msg": _extract_email,
        }

        if ext in extractors:
            text = extractors[ext](path)
        else:
            raise UnsupportedFileType(
                f"File type '{ext}' is not supported for text extraction. "
                "Supported: pdf, docx, txt, csv, json, eml, msg, mp4, m4a, mp3, wav."
            )

    # Normalise: collapse 3+ blank lines → 2, strip leading/trailing whitespace
    import re
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    return text


# ---------------------------------------------------------------------------
# Garbled-text detection (for PDF OCR fallback)
# ---------------------------------------------------------------------------

_COMMON_WORDS = {
    "the", "be", "to", "of", "and", "a", "in", "that", "have", "i",
    "it", "for", "not", "on", "with", "he", "as", "you", "do", "at",
    "this", "but", "his", "by", "from", "they", "we", "say", "her",
    "she", "or", "an", "will", "my", "one", "all", "would", "there",
    "their", "what", "so", "up", "out", "if", "about", "who", "get",
    "which", "go", "me", "when", "make", "can", "like", "time", "no",
    "just", "him", "know", "take", "people", "into", "year", "your",
    "good", "some", "could", "them", "see", "other", "than", "then",
    "now", "look", "only", "come", "its", "over", "think", "also",
    "back", "after", "use", "two", "how", "our", "work", "first",
    "well", "way", "even", "new", "want", "because", "any", "these",
    "give", "day", "most", "us", "is", "are", "was", "were", "been",
    "has", "had", "did", "does", "may", "shall", "should", "must",
    "form", "income", "tax", "total", "adjusted", "gross", "wages",
    "filing", "status", "exemptions", "deductions", "credits",
    "payment", "refund", "amount", "line", "schedule", "return",
    "federal", "state", "social", "security", "medicare", "withholding",
    "estimated", "balance", "due", "overpaid", "taxable", "net",
    "name", "address", "number", "date", "signature", "page",
    "department", "treasury", "internal", "revenue", "service",
}


def _is_garbled(text: str, threshold: float = 0.15) -> bool:
    """Detect whether extracted PDF text is garbled.

    Checks for:
    - Too short or too few words
    - Reversed words (even 1 means the PDF uses a custom font encoding
      that pdfplumber can't decode — the entire page is unreliable)
    - CID references like ``(cid:123)`` — unmapped glyphs
    - Low ratio of recognised English words
    """
    import re as _re

    if not text or len(text.strip()) < 20:
        return True

    words = _re.findall(r"[A-Za-z]{2,}", text)
    if len(words) < 5:
        return True

    # Check for CID references — unmapped glyphs, clear sign of encoding failure
    cid_count = len(_re.findall(r"\(cid:\d+\)", text))
    if cid_count >= 3:
        logger.info("Garbled detection: found %d (cid:N) references", cid_count)
        return True

    # Check for reversed words (strong signal — e.g. "mroF" for "Form").
    # Even ONE reversed common word means the PDF uses a custom font encoding
    # that pdfplumber can't reliably decode.  The page may look partially
    # readable but critical data (line numbers, dollar amounts) is often
    # missing or scrambled.
    reversed_hits = 0
    for w in words:
        if w.lower()[::-1] in _COMMON_WORDS and w.lower() not in _COMMON_WORDS:
            reversed_hits += 1
    if reversed_hits >= 1:
        logger.info("Garbled detection: found %d reversed words", reversed_hits)
        return True

    # Check ratio of recognised words
    recognised = sum(1 for w in words if w.lower() in _COMMON_WORDS)
    ratio = recognised / len(words)
    if ratio < threshold:
        logger.info(
            "Garbled detection: only %.1f%% recognised words (%d/%d)",
            ratio * 100, recognised, len(words),
        )
        return True

    return False


def _extract_pdf_ocr(path: Path) -> str:
    """Extract text from PDF using Tesseract OCR via pdf2image.

    Processes pages one at a time to avoid loading all images into memory,
    which can cause OOM crashes on large PDFs (e.g. 30-page IRS returns).
    """
    import gc
    from pdf2image import convert_from_path, pdfinfo_from_path
    import pytesseract

    logger.info(f"Running OCR extraction on {path.name}")

    info = pdfinfo_from_path(str(path))
    total_pages = info["Pages"]
    logger.info(f"OCR extraction: {path.name} has {total_pages} pages")

    pages: list[str] = []
    for page_num in range(1, total_pages + 1):
        try:
            images = convert_from_path(
                str(path),
                first_page=page_num,
                last_page=page_num,
                dpi=150,  # Higher DPI for accurate financial text extraction
            )
            image = images[0]

            page_text = pytesseract.image_to_string(image, config="--psm 6")
            if page_text and page_text.strip():
                pages.append(page_text.strip())
            logger.debug(f"OCR page {page_num}/{total_pages}: {len(page_text)} chars")

            image.close()
            del image, images
            gc.collect()
        except Exception as e:
            logger.error(f"OCR failed for page {page_num} of {path.name}: {e}")
            continue

    full_text = "\n\n".join(pages)
    logger.info(f"OCR complete: {total_pages} pages, {len(full_text)} chars")
    return full_text


# ---------------------------------------------------------------------------
# Per-format extractors
# ---------------------------------------------------------------------------


def _ocr_single_page(path: Path, page_num: int) -> str:
    """OCR a single page (1-indexed) via Tesseract. Returns extracted text or ''."""
    import gc
    from pdf2image import convert_from_path
    import pytesseract

    try:
        images = convert_from_path(
            str(path),
            first_page=page_num,
            last_page=page_num,
            dpi=150,
        )
        image = images[0]
        text = pytesseract.image_to_string(image, config="--psm 6")
        image.close()
        del image, images
        gc.collect()
        return (text or "").strip()
    except Exception as e:
        logger.error("OCR failed for page %d of %s: %s", page_num, path.name, e)
        return ""


def _extract_pdf(path: Path) -> str:
    """Extract text from PDF with per-page OCR fallback for garbled pages.

    Each page is checked individually for garbled text.  Pages where
    pdfplumber produces garbled output (common on Form 1040 from certain
    tax-software renderers) are re-extracted via Tesseract OCR, while
    clean pages keep their pdfplumber text.

    Every page is prefixed with a ``[Page N]`` marker so downstream
    chunking preserves page context for the LLM.
    """
    try:
        import pdfplumber
    except ImportError as exc:
        raise ExtractionError("pdfplumber is not installed") from exc

    # Step 1: Extract every page with pdfplumber
    pdfplumber_pages: list[tuple[int, str]] = []  # (1-indexed page_num, text)
    try:
        with pdfplumber.open(str(path)) as pdf:
            for i, page in enumerate(pdf.pages, start=1):
                text = page.extract_text()
                pdfplumber_pages.append((i, (text or "").strip()))
    except Exception as exc:
        raise ExtractionError(f"PDF extraction failed: {exc}") from exc

    if not pdfplumber_pages:
        raise ExtractionError(f"PDF has no pages: {path.name}")

    # Step 2: Per-page quality check — OCR only the garbled pages
    final_pages: list[str] = []
    ocr_count = 0

    for page_num, page_text in pdfplumber_pages:
        if _is_garbled(page_text):
            # This page is garbled — try OCR
            logger.info(
                "PDF extraction: page %d of %s is garbled, attempting OCR",
                page_num, path.name,
            )
            ocr_text = _ocr_single_page(path, page_num)
            if ocr_text and len(ocr_text) > len(page_text):
                page_text = ocr_text
                ocr_count += 1
            elif not page_text and ocr_text:
                page_text = ocr_text
                ocr_count += 1
            # else: keep pdfplumber text (both are poor; pdfplumber is at least consistent)

        if page_text:
            final_pages.append(f"[Page {page_num}]\n{page_text}")

    if not final_pages:
        # All pages empty — try full OCR as last resort
        logger.warning(
            "PDF extraction: all pdfplumber pages empty for %s, trying full OCR",
            path.name,
        )
        try:
            ocr_text = _extract_pdf_ocr(path)
            if ocr_text and ocr_text.strip():
                return ocr_text
        except Exception as e:
            logger.error("Full OCR fallback also failed: %s", e)
        raise ExtractionError(f"No text could be extracted from {path.name}")

    logger.info(
        "PDF extraction: %s — %d pages, %d needed OCR",
        path.name, len(pdfplumber_pages), ocr_count,
    )
    return "\n\n".join(final_pages)


def _extract_docx(path: Path) -> str:
    try:
        from docx import Document as DocxDocument
    except ImportError as exc:
        raise ExtractionError("python-docx is not installed") from exc

    try:
        doc = DocxDocument(str(path))
    except Exception as exc:
        raise ExtractionError(f"DOCX extraction failed: {exc}") from exc

    parts: list[str] = []

    # Paragraphs
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            parts.append(t)

    # Tables (row-per-line, cells tab-separated)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append("\t".join(cells))

    return "\n\n".join(parts)


def _extract_txt(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8", errors="replace")
    except Exception as exc:
        raise ExtractionError(f"TXT extraction failed: {exc}") from exc


def _extract_csv(path: Path) -> str:
    try:
        content = path.read_text(encoding="utf-8", errors="replace")
        reader = csv.reader(io.StringIO(content))
        rows = [", ".join(cell.strip() for cell in row) for row in reader if any(row)]
        return "\n".join(rows)
    except Exception as exc:
        raise ExtractionError(f"CSV extraction failed: {exc}") from exc


def _extract_fathom_json(path: Path) -> str:
    """
    Extract text from a Fathom meeting transcript JSON export.

    Produces a clean, readable document in this order:
        Meeting: <title>
        Date: <date>
        Attendees: <names>

        Summary:
        <summary>

        Transcript:
        [timestamp] Speaker: text
        ...

        Action Items:
        - Owner: description
        ...

    Field names are tried with common Fathom variants so the extractor
    stays robust across Fathom export versions.
    """
    import json as _json

    try:
        data = _json.loads(path.read_text(encoding="utf-8", errors="replace"))
    except _json.JSONDecodeError as exc:
        raise ExtractionError(f"JSON parsing failed: {exc}") from exc

    if not isinstance(data, dict):
        raise ExtractionError(
            f"Expected a JSON object at the top level, got {type(data).__name__}"
        )

    parts: list[str] = []

    # ------------------------------------------------------------------
    # 1. Meeting metadata header
    # ------------------------------------------------------------------
    header: list[str] = []

    title = (
        data.get("title")
        or data.get("name")
        or data.get("meeting_title")
        or data.get("meeting_name")
    )
    if title:
        header.append(f"Meeting: {title}")

    date_val = (
        data.get("date")
        or data.get("meeting_date")
        or data.get("call_date")
        or data.get("start_time")
        or data.get("created_at")
    )
    if date_val:
        header.append(f"Date: {date_val}")

    attendees = (
        data.get("attendees")
        or data.get("participants")
        or data.get("speakers")
        or []
    )
    if isinstance(attendees, list) and attendees:
        names: list[str] = []
        for a in attendees:
            if isinstance(a, dict):
                name = (
                    a.get("name") or a.get("full_name") or a.get("display_name")
                )
                if name:
                    names.append(str(name))
            elif isinstance(a, str) and a:
                names.append(a)
        if names:
            header.append(f"Attendees: {', '.join(names)}")
    elif isinstance(attendees, str) and attendees:
        header.append(f"Attendees: {attendees}")

    duration = data.get("duration") or data.get("duration_seconds")
    if duration:
        header.append(f"Duration: {duration}")

    if header:
        parts.append("\n".join(header))

    # ------------------------------------------------------------------
    # 2. Summary / overview
    # ------------------------------------------------------------------
    summary = (
        data.get("summary")
        or data.get("overview")
        or data.get("meeting_summary")
        or data.get("description")
    )
    if isinstance(summary, str) and summary.strip():
        parts.append(f"Summary:\n{summary.strip()}")

    # ------------------------------------------------------------------
    # 3. Transcript entries
    # ------------------------------------------------------------------
    transcript = (
        data.get("transcript")
        or data.get("transcript_entries")
        or data.get("entries")
        or data.get("utterances")
        or data.get("segments")
    )
    if transcript:
        if isinstance(transcript, list):
            lines: list[str] = []
            for entry in transcript:
                if isinstance(entry, dict):
                    speaker = (
                        entry.get("speaker")
                        or entry.get("speaker_name")
                        or entry.get("name")
                        or ""
                    )
                    text = (
                        entry.get("text")
                        or entry.get("content")
                        or entry.get("transcript")
                        or ""
                    )
                    timestamp = (
                        entry.get("start")
                        or entry.get("start_time")
                        or entry.get("timestamp")
                        or ""
                    )
                    text = str(text).strip()
                    if not text:
                        continue
                    if speaker:
                        prefix = f"[{timestamp}] {speaker}" if timestamp else speaker
                        lines.append(f"{prefix}: {text}")
                    else:
                        lines.append(text)
                elif isinstance(entry, str) and entry.strip():
                    lines.append(entry.strip())
            if lines:
                parts.append("Transcript:\n" + "\n".join(lines))
        elif isinstance(transcript, str) and transcript.strip():
            parts.append(f"Transcript:\n{transcript.strip()}")

    # ------------------------------------------------------------------
    # 4. Action items
    # ------------------------------------------------------------------
    action_items = (
        data.get("action_items")
        or data.get("tasks")
        or data.get("follow_ups")
        or data.get("next_steps")
        or data.get("todos")
    )
    if action_items:
        if isinstance(action_items, list):
            items: list[str] = []
            for item in action_items:
                if isinstance(item, dict):
                    owner = (
                        item.get("owner")
                        or item.get("assignee")
                        or item.get("assigned_to")
                        or ""
                    )
                    desc = (
                        item.get("description")
                        or item.get("text")
                        or item.get("title")
                        or item.get("task")
                        or ""
                    )
                    desc = str(desc).strip()
                    if not desc:
                        continue
                    items.append(f"- {owner}: {desc}" if owner else f"- {desc}")
                elif isinstance(item, str) and item.strip():
                    items.append(f"- {item.strip()}")
            if items:
                parts.append("Action Items:\n" + "\n".join(items))
        elif isinstance(action_items, str) and action_items.strip():
            parts.append(f"Action Items:\n{action_items.strip()}")

    if not parts:
        raise ExtractionError(
            "No readable content found in JSON. "
            "Expected Fathom transcript fields: title, date, attendees, "
            "transcript, action_items."
        )

    return "\n\n".join(parts)


def _extract_email(path: Path) -> str:
    """Extract text from .eml or .msg email files via email_extractor."""
    try:
        from app.services.email_extractor import extract_email_text
        return extract_email_text(str(path))
    except ValueError as exc:
        raise ExtractionError(f"Email extraction failed: {exc}") from exc
    except Exception as exc:
        raise ExtractionError(f"Email extraction failed: {exc}") from exc
